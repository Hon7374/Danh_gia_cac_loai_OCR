# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path.cwd()
REPORT_DIR = ROOT / "reports" / "benchmark_technical_report"
ASSETS = REPORT_DIR / "assets"
MARKDOWN_PATH = REPORT_DIR / "OCR_BENCHMARK_TECHNICAL_REPORT.md"
DOCX_PATH = REPORT_DIR / "OCR_BENCHMARK_TECHNICAL_REPORT.docx"
DOCX_FALLBACK_PATH = REPORT_DIR / "OCR_BENCHMARK_TECHNICAL_REPORT_UPDATED.docx"


def font_path(name: str = "arial.ttf") -> str:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/tahoma.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return ""


def pil_font(name: str, size: int):
    path = font_path(name)
    if path:
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#1F4E79", width: int = 4) -> None:
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    left = (x2 - 18 * ux - 9 * uy, y2 - 18 * uy + 9 * ux)
    right = (x2 - 18 * ux + 9 * uy, y2 - 18 * uy - 9 * ux)
    draw.polygon([end, left, right], fill=color)


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body: str = "", fill: str = "#EAF2FF") -> None:
    title_font = pil_font("arialbd.ttf", 28)
    body_font = pil_font("arial.ttf", 20)
    draw.rounded_rectangle(box, radius=22, fill=fill, outline="#9CB9D8", width=3)
    x1, y1, x2, y2 = box
    draw.text((x1 + 22, y1 + 22), title, fill="#003B73", font=title_font)
    if body:
        y = y1 + 62
        for line in textwrap.wrap(body, width=28):
            draw.text((x1 + 22, y), line, fill="#24445F", font=body_font)
            y += 28


def create_architecture_image() -> Path:
    path = ASSETS / "architecture_workflow.png"
    image = Image.new("RGB", (1900, 1050), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = pil_font("arialbd.ttf", 42)
    draw.text((70, 45), "Kiến trúc hệ thống OCR + chữ ký số", fill="#0B1F33", font=title_font)

    boxes = [
        ((70, 150, 360, 285), "Người dùng", "Upload PDF hoặc scan", "#EAF2FF"),
        ((445, 150, 735, 285), "Tiền xử lý", "Tách trang, raw branch, OpenCV", "#EAFBF3"),
        ((820, 150, 1110, 285), "OCR benchmark", "Tesseract, EasyOCR, PaddleOCR, PaddleOCR-VL", "#FFF4E6"),
        ((1195, 150, 1485, 285), "Chọn OCR tốt nhất", "So runtime, text length, confidence, quality", "#F6EDFF"),
        ((1560, 150, 1850, 285), "LayoutLMv3", "Document understanding, metadata extraction", "#EAF2FF"),
        ((445, 430, 735, 565), "Metadata", "Số ký hiệu, ngày, cơ quan, trích yếu", "#EAFBF3"),
        ((820, 430, 1110, 565), "Chữ ký số", "Kiểm tra chứng thư, toàn vẹn file", "#FFF4E6"),
        ((1195, 430, 1485, 565), "Kho dữ liệu", "DB, file gốc, OCR output, metadata", "#F6EDFF"),
        ((1560, 430, 1850, 565), "Dashboard", "Tra cứu, benchmark, đánh giá", "#EAF2FF"),
    ]
    for box, title, body, fill in boxes:
        rounded_box(draw, box, title, body, fill)

    arrows = [
        ((360, 218), (445, 218)),
        ((735, 218), (820, 218)),
        ((1110, 218), (1195, 218)),
        ((1485, 218), (1560, 218)),
        ((1705, 285), (1705, 430)),
        ((1560, 498), (1485, 498)),
        ((1195, 498), (1110, 498)),
        ((820, 498), (735, 498)),
    ]
    for start, end in arrows:
        draw_arrow(draw, start, end)

    note_font = pil_font("arial.ttf", 22)
    draw.rounded_rectangle((70, 720, 1850, 930), radius=20, fill="#F8FAFC", outline="#D8E2EF", width=2)
    note = (
        "Luồng chính: người dùng tải công văn lên, hệ thống tách trang và tạo nhánh raw/OpenCV, "
        "chạy benchmark OCR, chọn text phù hợp nhất, dùng LayoutLMv3 để hiểu cấu trúc tài liệu, "
        "sau đó kiểm tra chữ ký số và lưu hồ sơ vào kho dữ liệu."
    )
    y = 760
    for line in textwrap.wrap(note, width=130):
        draw.text((105, y), line, fill="#0B1F33", font=note_font)
        y += 34

    image.save(path)
    return path


def create_sequence_image() -> Path:
    path = ASSETS / "sequence_diagram.png"
    image = Image.new("RGB", (1900, 1150), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = pil_font("arialbd.ttf", 42)
    label_font = pil_font("arialbd.ttf", 24)
    msg_font = pil_font("arial.ttf", 21)
    draw.text((70, 45), "Sequence xử lý công văn điện tử", fill="#0B1F33", font=title_font)

    participants = [
        ("User", 150),
        ("System", 470),
        ("OCR Engines", 790),
        ("LayoutLMv3", 1110),
        ("Signature Verification", 1430),
        ("Database", 1740),
    ]
    for label, x in participants:
        draw.rounded_rectangle((x - 115, 130, x + 115, 195), radius=16, fill="#EAF2FF", outline="#9CB9D8", width=2)
        bbox = draw.textbbox((0, 0), label, font=label_font)
        draw.text((x - (bbox[2] - bbox[0]) / 2, 148), label, fill="#003B73", font=label_font)
        draw.line([(x, 195), (x, 1040)], fill="#B5C7D8", width=3)

    messages = [
        ("User", "System", "Upload PDF/scan công văn"),
        ("System", "System", "Tách trang ảnh và sinh nhánh OpenCV"),
        ("System", "OCR Engines", "Chạy OCR raw/OpenCV"),
        ("OCR Engines", "System", "Trả text, runtime, confidence, quality"),
        ("System", "LayoutLMv3", "Gửi OCR output + layout/image"),
        ("LayoutLMv3", "System", "Trả metadata và nhãn field"),
        ("System", "Signature Verification", "Kiểm tra chữ ký số/PAdES"),
        ("Signature Verification", "System", "Trả trạng thái chữ ký"),
        ("System", "Database", "Lưu file, OCR, metadata, kết quả đánh giá"),
        ("User", "System", "Tra cứu dashboard/report"),
        ("System", "User", "Hiển thị benchmark và hồ sơ"),
    ]
    x_map = dict(participants)
    y = 260
    for src, dst, text in messages:
        x1, x2 = x_map[src], x_map[dst]
        if src == dst:
            draw.line([(x1, y), (x1 + 105, y), (x1 + 105, y + 38), (x1, y + 38)], fill="#1F4E79", width=3)
            draw_arrow(draw, (x1 + 105, y + 38), (x1 + 5, y + 38), width=3)
            tx = x1 + 120
        else:
            draw_arrow(draw, (x1, y), (x2, y), width=3)
            tx = min(x1, x2) + 25
        draw.text((tx, y - 30), text, fill="#0B1F33", font=msg_font)
        y += 72

    image.save(path)
    return path


def set_run_font(run, name: str = "Times New Roman", size: int | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline_markdown(paragraph, text: str, size: int = 13) -> None:
    text = text.replace("  ", " ")
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=max(size - 1, 8))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clean_table_cell(text: str) -> str:
    text = text.strip()
    text = text.replace("<br>", "\n")
    text = text.replace("\\|", "|")
    text = re.sub(r"^\s*`(.*)`\s*$", r"\1", text)
    text = text.replace("**", "")
    return text


def split_table_row(line: str) -> list[str]:
    return [clean_table_cell(part) for part in line.strip().strip("|").split("|")]


def add_markdown_table(document: Document, table_lines: list[str]) -> None:
    rows = [split_table_row(line) for line in table_lines if not re.match(r"^\s*\|?\s*:?-{3,}", line.replace("|", " |"))]
    rows = [row for row in rows if row]
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    small = max(len(rows[0]) >= 7, False)
    font_size = 8 if small else 10

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markdown(paragraph, value, size=font_size)
            for run in paragraph.runs:
                if row_idx == 0:
                    run.bold = True
    document.add_paragraph()


def add_image(document: Document, image_path: Path, alt_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    filename = image_path.name
    if filename in {"input_raw_page1.png", "input_opencv_page1.png"}:
        width = Cm(13.5)
    else:
        width = Cm(16.2)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width)

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(f"Hình: {alt_text}")
    set_run_font(caption_run, size=11, italic=True)


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["No Spacing"]
    run = paragraph.add_run(code)
    set_run_font(run, name="Courier New", size=9)


def setup_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)

    for style_name, size, color in [
        ("Heading 1", 18, "003B73"),
        ("Heading 2", 15, "004A99"),
        ("Heading 3", 13, "0B1F33"),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def convert_markdown_to_docx() -> None:
    create_architecture_image()
    create_sequence_image()

    markdown = MARKDOWN_PATH.read_text(encoding="utf-8").splitlines()
    document = Document()
    setup_document(document)
    document.core_properties.title = "Báo cáo kỹ thuật benchmark OCR và trích xuất metadata công văn tiếng Việt"
    document.core_properties.subject = "Ứng dụng OCR kết hợp chữ ký số trong quản lý công văn điện tử tiếng Việt"

    idx = 0
    while idx < len(markdown):
        line = markdown[idx]
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            code_type = stripped.strip("`").strip()
            idx += 1
            while idx < len(markdown) and not markdown[idx].strip().startswith("```"):
                code_lines.append(markdown[idx])
                idx += 1
            idx += 1
            code = "\n".join(code_lines)
            if code_type == "mermaid" and "graph TD" in code:
                add_image(document, ASSETS / "architecture_workflow.png", "Kiến trúc hệ thống OCR kết hợp chữ ký số")
            elif code_type == "mermaid" and "sequenceDiagram" in code:
                add_image(document, ASSETS / "sequence_diagram.png", "Sequence xử lý công văn điện tử")
            else:
                add_code_block(document, code)
            continue

        if stripped.startswith("|"):
            table_lines = []
            while idx < len(markdown) and markdown[idx].strip().startswith("|"):
                table_lines.append(markdown[idx])
                idx += 1
            add_markdown_table(document, table_lines)
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            alt_text, rel_path = image_match.groups()
            add_image(document, REPORT_DIR / rel_path, alt_text)
            idx += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            paragraph = document.add_heading(level=level)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markdown(paragraph, text, size={1: 18, 2: 15, 3: 13}.get(level, 13))
            idx += 1
            continue

        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.space_after = Pt(6)
            add_inline_markdown(paragraph, stripped.lstrip("> ").strip(), size=12)
            for run in paragraph.runs:
                run.italic = True
            idx += 1
            continue

        ordered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, ordered.group(2), size=13)
            idx += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:], size=13)
            idx += 1
            continue

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(0.6)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(6)
        add_inline_markdown(paragraph, stripped, size=13)
        idx += 1

    try:
        document.save(DOCX_PATH)
        print(f"Wrote {DOCX_PATH}")
    except PermissionError:
        document.save(DOCX_FALLBACK_PATH)
        print(f"Wrote {DOCX_FALLBACK_PATH}")


if __name__ == "__main__":
    convert_markdown_to_docx()
